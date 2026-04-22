import asyncio
import io
import logging
import os
import tempfile
import time
from typing import Any
from uuid import UUID, uuid4

from qdrant_client.models import PointStruct

from ..config import settings
from ..db import SessionLocal
from ..models import Chunk, FileRecord
from ..progress_broker import broker
from . import bifrost, qdrant, s3

log = logging.getLogger("task.ingest")


def _prefix(task: str, file_id: UUID, filename: str | None = None) -> str:
    """Every log line for a background task starts with one of these so you
    can grep a single task end-to-end."""
    tag = str(file_id)[:8]
    if filename:
        return f"[{task} {tag} {filename!r}]"
    return f"[{task} {tag}]"

_semaphore: asyncio.Semaphore | None = None
# Serializes Docling invocations. Only PPTX still routes through Docling
# (with OCR, because slide content is often rasterised); every other format
# now uses a lightweight native extractor (pymupdf / python-docx /
# openpyxl). Docling's peak RSS is ~1.5–2 GB while converting, so two at
# once would tip over an 8 GB box — keep the lock even though only one
# format uses it, in case INGEST_CONCURRENT_FILES > 1 and two PPTX files
# land in flight together.
_docling_lock_obj: asyncio.Lock | None = None


def _sem() -> asyncio.Semaphore:
    global _semaphore
    if _semaphore is None:
        _semaphore = asyncio.Semaphore(settings.INGEST_CONCURRENT_FILES)
    return _semaphore


def _docling_lock() -> asyncio.Lock:
    global _docling_lock_obj
    if _docling_lock_obj is None:
        _docling_lock_obj = asyncio.Lock()
    return _docling_lock_obj


# Docling converter for PPTX. Building one allocates ~1.5–2 GB of PyTorch
# weights (layout + OCR models), so we load it lazily at first PPTX ingest
# and keep one process-global singleton. If PPTX uploads are rare this
# stays unloaded and the container footprint remains small. Every other
# format uses lightweight native extractors and never touches Docling.
_DOCLING_OCR = None  # type: ignore[var-annotated]


def _docling_ocr_converter():
    """Lazy singleton loaded on first PPTX ingest. OCR is enabled because
    slide decks often have text rendered as images (diagrams, exported
    screenshots) that native text extraction would miss."""
    global _DOCLING_OCR
    if _DOCLING_OCR is not None:
        return _DOCLING_OCR
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    opts = PdfPipelineOptions(
        do_ocr=True,
        do_table_structure=False,
        generate_page_images=False,
    )
    _DOCLING_OCR = DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
    )
    log.warning(
        "loaded Docling-OCR singleton for PPTX (adds ~1.5-2 GB RSS, "
        "resident for process lifetime)"
    )
    return _DOCLING_OCR


# ---------- sync DB helpers (called via to_thread) ----------
def _load_file_row(file_id: UUID) -> dict[str, Any]:
    db = SessionLocal()
    try:
        row = db.get(FileRecord, file_id)
        if row is None:
            raise RuntimeError(f"file {file_id} not found")
        return {"s3_key": row.s3_key, "filename": row.filename}
    finally:
        db.close()


def _update_status(file_id: UUID, **kwargs) -> None:
    db = SessionLocal()
    try:
        row = db.get(FileRecord, file_id)
        if row is None:
            return
        for k, v in kwargs.items():
            setattr(row, k, v)
        db.commit()
    finally:
        db.close()


def _insert_chunk_rows(file_id: UUID, records: list[dict]) -> None:
    db = SessionLocal()
    try:
        db.bulk_insert_mappings(
            Chunk,
            [
                {
                    "id": uuid4(),
                    "file_id": file_id,
                    "qdrant_point_id": r["point_id"],
                    "page": r.get("page"),
                    "chunk_text": r["text"],
                    "element_type": None,
                }
                for r in records
            ],
        )
        db.commit()
    finally:
        db.close()


# ---------- parsing + chunking ----------
def _download(s3_key: str, prefix: str) -> bytes:
    t0 = time.perf_counter()
    log.info("%s downloading from S3 key=%s", prefix, s3_key)
    resp = s3.client().get_object(Bucket=settings.S3_BUCKET, Key=s3_key)
    data = resp["Body"].read()
    log.info(
        "%s downloaded %d bytes in %.2fs",
        prefix,
        len(data),
        time.perf_counter() - t0,
    )
    return data


def _pymupdf_open(data: bytes, prefix: str):
    """Open a PDF via pymupdf. Returns (doc, page_count) or None if the
    file can't be opened. The caller owns `doc` and must close it.

    We also log chars/page over a small sample so that scanned PDFs
    (which will produce zero extractable text and fail downstream) are
    obvious in the logs without us having to eagerly read every page."""
    import fitz  # pymupdf

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as e:  # noqa: BLE001
        log.info("%s pymupdf could not open file (%s)", prefix, e)
        return None
    try:
        page_count = doc.page_count
        sample = min(page_count, 8)
        total = 0
        for i in range(sample):
            total += len(doc[i].get_text("text") or "")
        avg = total / max(1, sample)
        if avg < settings.PYMUPDF_MIN_CHARS_PER_PAGE:
            log.warning(
                "%s pymupdf sampled %.1f chars/page over %d of %d pages -- "
                "looks scanned; ingest will produce zero chunks and fail",
                prefix,
                avg,
                sample,
                page_count,
            )
        else:
            log.info(
                "%s pymupdf probe %.0f chars/page over %d sample pages (total %d pages)",
                prefix,
                avg,
                sample,
                page_count,
            )
        return doc, page_count
    except Exception:
        doc.close()
        raise


def _extract_docx_markdown(data: bytes, prefix: str) -> str:
    """Extract paragraphs + tables from a DOCX without loading Docling.
    Images/embedded objects are ignored (no OCR)."""
    from docx import Document  # python-docx

    t0 = time.perf_counter()
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    md = "\n\n".join(parts)
    log.info(
        "%s python-docx extracted %d chars in %.2fs",
        prefix,
        len(md),
        time.perf_counter() - t0,
    )
    return md


def _extract_xlsx_markdown(data: bytes, prefix: str) -> str:
    """Extract cell values from every sheet. read_only=True streams rows
    instead of building the full object model, which matters for big
    spreadsheets. data_only=True resolves formula results rather than the
    formula strings."""
    from openpyxl import load_workbook

    t0 = time.perf_counter()
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"## {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
    finally:
        wb.close()
    md = "\n\n".join(parts)
    log.info(
        "%s openpyxl extracted %d chars in %.2fs",
        prefix,
        len(md),
        time.perf_counter() - t0,
    )
    return md


def _extract_txt(data: bytes, prefix: str) -> str:
    """Plain-text files: best-effort decode. errors='replace' keeps the
    pipeline moving on files with stray bytes rather than failing."""
    t0 = time.perf_counter()
    md = data.decode("utf-8", errors="replace")
    log.info(
        "%s txt decoded %d chars in %.2fs",
        prefix,
        len(md),
        time.perf_counter() - t0,
    )
    return md


def _extract_pdf_slab(doc, start: int, end: int) -> str:
    """Pull text from pages [start, end) out of an already-open fitz doc and
    return a single markdown slab. Called via to_thread so we don't stall the
    event loop while pymupdf churns through pages."""
    parts: list[str] = []
    for i in range(start, end):
        t = (doc[i].get_text("text") or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _parse_pptx_with_docling(data: bytes, filename: str, prefix: str) -> str:
    """Blocking -- call via asyncio.to_thread. PPTX is the only format that
    still goes through Docling (with OCR enabled), because slide content is
    frequently rendered as images rather than live text shapes. Uses a
    module-level singleton so Docling models load at most once per process."""
    suffix = os.path.splitext(filename)[1] or ""
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tf:
        tf.write(data)
        tmp_path = tf.name
    try:
        log.info(
            "%s parsing PPTX with Docling (OCR enabled, first run loads models)...",
            prefix,
        )
        t0 = time.perf_counter()
        converter = _docling_ocr_converter()
        result = converter.convert(tmp_path)
        md = result.document.export_to_markdown()
        log.info(
            "%s Docling parse done in %.2fs (%d chars of markdown)",
            prefix,
            time.perf_counter() - t0,
            len(md),
        )
        return md
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _ext(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def _chunk_markdown(text: str, target_chars: int = 2000, overlap_chars: int = 200) -> list[str]:
    """Paragraph-aware sliding window. Target ~500 tokens (≈2000 chars)."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buf = ""
    for p in paragraphs:
        if not buf:
            buf = p
            continue
        if len(buf) + 2 + len(p) <= target_chars:
            buf = buf + "\n\n" + p
            continue
        chunks.append(buf)
        tail = buf[-overlap_chars:] if overlap_chars else ""
        buf = (tail + "\n\n" + p) if tail else p
    if buf:
        chunks.append(buf)
    return chunks


# ---------- orchestrator ----------
async def run_ingest(file_id: UUID) -> None:
    """Entry point scheduled via BackgroundTasks. Guarded by a global
    semaphore so at most INGEST_CONCURRENT_FILES run at the same time."""
    try:
        meta = await asyncio.to_thread(_load_file_row, file_id)
    except Exception:
        log.exception("[ingest %s] could not load file row -- aborting", str(file_id)[:8])
        return

    filename = meta["filename"]
    s3_key = meta["s3_key"]
    prefix = _prefix("ingest", file_id, filename)

    log.info("%s ====== starting ingest task ======", prefix)
    log.info("%s file_id=%s s3_key=%s", prefix, file_id, s3_key)

    sem = _sem()
    if sem.locked():
        log.info(
            "%s all %d concurrency slots busy, waiting for a slot...",
            prefix,
            settings.INGEST_CONCURRENT_FILES,
        )
    async with sem:
        log.info("%s acquired concurrency slot", prefix)
        task_t0 = time.perf_counter()
        try:
            await _do_ingest(file_id, filename, s3_key, prefix)
            log.info(
                "%s ====== done in %.2fs -- status=ready ======",
                prefix,
                time.perf_counter() - task_t0,
            )
        except Exception as e:  # noqa: BLE001 -- top-level task guard
            log.exception(
                "%s FAILED after %.2fs -- %s",
                prefix,
                time.perf_counter() - task_t0,
                e,
            )
            await asyncio.to_thread(
                _update_status, file_id, status="failed", error_message=str(e)[:500]
            )
            broker.publish(str(file_id), {"status": "failed", "error": str(e)[:200]})


PDF_STREAM_PAGE_SLAB = 20


async def _process_markdown_slab(
    file_id: UUID,
    filename: str,
    prefix: str,
    markdown: str,
    chunk_index_base: int,
) -> int:
    """Chunk one markdown slab, then for each embedding batch: embed, upsert
    to Qdrant, and persist chunk rows to Postgres — releasing each batch's
    buffers before moving to the next. Returns the number of chunks that were
    produced and stored for this slab.

    Per-batch persistence (vs. the old end-of-file accumulation) is what
    keeps peak RAM bounded on very large documents: at any instant we hold
    one batch's worth of text + vectors + point structs, not the whole
    document's."""
    chunks = _chunk_markdown(markdown)
    if not chunks:
        return 0
    batch_size = settings.EMBED_MAX_BATCH_ITEMS
    slab_total = len(chunks)
    batch_total = (slab_total + batch_size - 1) // batch_size

    for batch_idx, start in enumerate(range(0, slab_total, batch_size), 1):
        batch = chunks[start : start + batch_size]
        log.info(
            "%s embedding slab batch %d/%d (%d items)...",
            prefix,
            batch_idx,
            batch_total,
            len(batch),
        )
        t_embed = time.perf_counter()
        vectors = await bifrost.embed_texts(batch)
        log.info(
            "%s  embed done in %.2fs -- upserting to Qdrant",
            prefix,
            time.perf_counter() - t_embed,
        )
        points = []
        records: list[dict] = []
        for j, (text, vec) in enumerate(zip(batch, vectors)):
            point_id = str(uuid4())
            points.append(
                PointStruct(
                    id=point_id,
                    vector=vec,
                    payload={
                        "file_id": str(file_id),
                        "filename": filename,
                        "chunk_index": chunk_index_base + start + j,
                        "chunk_text": text,
                    },
                )
            )
            records.append({"point_id": point_id, "text": text})
        t_upsert = time.perf_counter()
        await qdrant.client().upsert(
            collection_name=settings.QDRANT_COLLECTION, points=points
        )
        log.info(
            "%s  upsert done in %.2fs (%d points)",
            prefix,
            time.perf_counter() - t_upsert,
            len(points),
        )
        # Persist chunk rows for this batch only — nothing accumulates across
        # batches or slabs.
        await asyncio.to_thread(_insert_chunk_rows, file_id, records)
        del points, records, vectors, batch

    return slab_total


async def _do_ingest(file_id: UUID, filename: str, s3_key: str, prefix: str) -> None:
    # 1. parse + route
    await asyncio.to_thread(
        _update_status, file_id, status="parsing", stage_current=0, stage_total=0
    )
    broker.publish(str(file_id), {"status": "parsing", "filename": filename})

    data = await asyncio.to_thread(_download, s3_key, prefix)

    # Flip to 'embedding' before any slab is processed; with streaming there
    # is no distinct 'chunking' phase visible to the user — chunking happens
    # interleaved with embedding slab-by-slab.
    await asyncio.to_thread(_update_status, file_id, status="embedding")
    broker.publish(
        str(file_id), {"status": "embedding", "stage_current": 0, "stage_total": 0}
    )

    total_chunks = 0

    ext = _ext(filename)

    if ext == ".pdf":
        probe = await asyncio.to_thread(_pymupdf_open, data, prefix)
        if probe is None:
            raise RuntimeError("pymupdf could not open PDF")
        doc, page_count = probe
        # fitz copies the bytes on open, so the original buffer is no longer
        # needed — free it before starting the per-slab loop.
        del data
        try:
            for start_page in range(0, page_count, PDF_STREAM_PAGE_SLAB):
                end_page = min(start_page + PDF_STREAM_PAGE_SLAB, page_count)
                t_slab = time.perf_counter()
                slab = await asyncio.to_thread(
                    _extract_pdf_slab, doc, start_page, end_page
                )
                log.info(
                    "%s slab pages %d-%d extracted in %.2fs (%d chars)",
                    prefix,
                    start_page + 1,
                    end_page,
                    time.perf_counter() - t_slab,
                    len(slab),
                )
                produced = await _process_markdown_slab(
                    file_id, filename, prefix, slab, total_chunks
                )
                total_chunks += produced
                del slab
                # Publish incremental progress. stage_total grows as we go;
                # the UI just reflects whatever the latest snapshot says.
                await asyncio.to_thread(
                    _update_status,
                    file_id,
                    stage_current=total_chunks,
                    stage_total=total_chunks,
                )
                broker.publish(
                    str(file_id),
                    {
                        "status": "embedding",
                        "stage_current": total_chunks,
                        "stage_total": total_chunks,
                    },
                )
        finally:
            doc.close()
    else:
        # Non-PDF: dispatch to the right extractor. DOCX/XLSX/TXT use
        # lightweight native libraries (no Docling, no OCR, low RAM).
        # PPTX still goes through Docling-with-OCR because slides often
        # have text rendered as images that native parsers would miss.
        if ext == ".docx":
            md = await asyncio.to_thread(_extract_docx_markdown, data, prefix)
        elif ext == ".xlsx":
            md = await asyncio.to_thread(_extract_xlsx_markdown, data, prefix)
        elif ext in (".txt", ".md"):
            md = await asyncio.to_thread(_extract_txt, data, prefix)
        elif ext == ".pptx":
            async with _docling_lock():
                md = await asyncio.to_thread(
                    _parse_pptx_with_docling, data, filename, prefix
                )
        else:
            raise RuntimeError(f"unsupported file extension: {ext or '(none)'}")
        del data

        total_chunks = await _process_markdown_slab(
            file_id, filename, prefix, md, 0
        )
        del md
        await asyncio.to_thread(
            _update_status,
            file_id,
            stage_current=total_chunks,
            stage_total=total_chunks,
        )
        broker.publish(
            str(file_id),
            {
                "status": "embedding",
                "stage_current": total_chunks,
                "stage_total": total_chunks,
            },
        )

    if total_chunks == 0:
        raise RuntimeError("document produced zero chunks")

    await asyncio.to_thread(_update_status, file_id, status="ready")
    broker.publish(
        str(file_id),
        {"status": "ready", "stage_current": total_chunks, "stage_total": total_chunks},
    )


# ---------- delete orchestrator ----------
def _fetch_s3_key(file_id: UUID) -> str | None:
    db = SessionLocal()
    try:
        row = db.get(FileRecord, file_id)
        return row.s3_key if row else None
    finally:
        db.close()


def _drop_row(file_id: UUID) -> None:
    db = SessionLocal()
    try:
        row = db.get(FileRecord, file_id)
        if row is None:
            return
        db.delete(row)
        db.commit()
    finally:
        db.close()


async def run_delete(file_id: UUID) -> None:
    """Background delete pipeline: Qdrant points → S3 object → DB row.
    Ordering matters -- we need the s3_key from the row, so DB goes last.
    On failure, records which stage blew up so the UI can surface it."""
    try:
        filename = (await asyncio.to_thread(_load_file_row, file_id))["filename"]
    except Exception:
        filename = None  # row may be missing on retry

    prefix = _prefix("delete", file_id, filename)
    log.info("%s ====== starting delete task ======", prefix)
    task_t0 = time.perf_counter()

    stage = "start"
    try:
        stage = "qdrant"
        log.info("%s stage=qdrant -- removing all points for file_id", prefix)
        t = time.perf_counter()
        await qdrant.delete_points_for_file(file_id)
        log.info("%s  qdrant done in %.2fs", prefix, time.perf_counter() - t)

        stage = "s3"
        s3_key = await asyncio.to_thread(_fetch_s3_key, file_id)
        if s3_key:
            log.info("%s stage=s3 -- removing object key=%s", prefix, s3_key)
            t = time.perf_counter()
            await asyncio.to_thread(s3.delete_object, s3_key)
            log.info("%s  s3 done in %.2fs", prefix, time.perf_counter() - t)
        else:
            log.info("%s stage=s3 -- no s3_key to delete (already gone)", prefix)

        stage = "db"
        log.info("%s stage=db -- dropping row (cascades to chunks)", prefix)
        t = time.perf_counter()
        await asyncio.to_thread(_drop_row, file_id)
        log.info("%s  db done in %.2fs", prefix, time.perf_counter() - t)

        broker.publish(str(file_id), {"status": "deleted"})
        log.info(
            "%s ====== done in %.2fs -- status=deleted ======",
            prefix,
            time.perf_counter() - task_t0,
        )
    except Exception as e:  # noqa: BLE001
        err = f"{stage}: {str(e)[:300]}"
        log.exception("%s FAILED at stage=%s after %.2fs", prefix, stage, time.perf_counter() - task_t0)
        await asyncio.to_thread(
            _update_status, file_id, status="delete_failed", error_message=err
        )
        broker.publish(str(file_id), {"status": "delete_failed", "error": err})


# ---------- dispatcher ----------
# A single app-level asyncio task owns ingest scheduling. The admin click on
# "Start Ingestion" flips staged rows -> queued and calls notify(); the
# dispatcher pulls queued rows FIFO and hands each to run_ingest (which has
# its own semaphore gating). This replaces FastAPI BackgroundTasks, which
# fired every upload's ingest simultaneously and caused the OOMs.
_dispatch_wake: asyncio.Event | None = None
_dispatch_task: asyncio.Task | None = None

log_disp = logging.getLogger("task.dispatcher")


def _claim_next_queued() -> UUID | None:
    """Atomically move one oldest 'queued' row to 'pending_ingest' and return
    its id. Returning None means no work waiting."""
    db = SessionLocal()
    try:
        row = (
            db.query(FileRecord)
            .filter(FileRecord.status == "queued")
            .order_by(FileRecord.created_at.asc())
            .with_for_update(skip_locked=True)
            .first()
        )
        if row is None:
            return None
        row.status = "pending_ingest"
        db.commit()
        fid = row.id
        broker.publish(str(fid), {"status": "pending_ingest", "filename": row.filename})
        return fid
    finally:
        db.close()


def _count_queued_or_running() -> int:
    db = SessionLocal()
    try:
        return (
            db.query(FileRecord)
            .filter(
                FileRecord.status.in_(
                    ["queued", "pending_ingest", "parsing", "chunking", "embedding"]
                )
            )
            .count()
        )
    finally:
        db.close()


def notify_dispatcher() -> None:
    """Poke the dispatcher to look for new queued rows. Safe to call from any
    sync context as long as the event loop is running."""
    if _dispatch_wake is not None:
        _dispatch_wake.set()


async def _dispatcher_loop() -> None:
    global _dispatch_wake
    _dispatch_wake = asyncio.Event()
    log_disp.info(
        "dispatcher started (concurrent_limit=%d)", settings.INGEST_CONCURRENT_FILES
    )

    # On boot, pick up anything left in 'queued' from a previous run.
    _dispatch_wake.set()

    running: set[asyncio.Task] = set()
    sem = _sem()

    while True:
        await _dispatch_wake.wait()
        _dispatch_wake.clear()

        while True:
            # Drain finished tasks so 'running' tracks only live work.
            running = {t for t in running if not t.done()}

            # Gate on semaphore availability. If all slots busy, stop pulling
            # -- we'll resume when a slot frees (via notify on completion).
            if sem.locked():
                break

            file_id = await asyncio.to_thread(_claim_next_queued)
            if file_id is None:
                break

            async def _run_and_notify(fid: UUID):
                try:
                    await run_ingest(fid)
                finally:
                    notify_dispatcher()

            t = asyncio.create_task(_run_and_notify(file_id))
            running.add(t)


def start_dispatcher() -> asyncio.Task:
    """Called from app lifespan. Idempotent."""
    global _dispatch_task
    if _dispatch_task is None or _dispatch_task.done():
        _dispatch_task = asyncio.create_task(_dispatcher_loop())
    return _dispatch_task


def stage_to_queued() -> int:
    """Flip every 'staged' row to 'queued' and kick the dispatcher. Returns
    the count of rows moved."""
    db = SessionLocal()
    try:
        rows = db.query(FileRecord).filter(FileRecord.status == "staged").all()
        for r in rows:
            r.status = "queued"
            r.error_message = None
            broker.publish(str(r.id), {"status": "queued", "filename": r.filename})
        db.commit()
        count = len(rows)
    finally:
        db.close()
    if count > 0:
        notify_dispatcher()
    return count
